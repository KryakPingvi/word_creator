from tkinter import *

from docx import Document
from docx.shared import Mm,Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph,ParagraphFormat
from docx.enum.text import WD_LINE_SPACING


def create_file(law_range,law_range_address,law_number_first,law_number_last,law_number_year,physics,physics_address,creditor,creditor_address,summ):
    #=============================================================================#
    #============================Создание файла===================================#

    doc = Document()
    #===============================Шапка=========================================#

    #
    # задаем стиль текста по умолчанию
    style = doc.styles['Normal']
    # название шрифта
    style.font.name = 'Times New Roman'
    # размер шрифта
    style.font.size = Pt(12)



    p = doc.add_paragraph()
    p.add_run(f'В Арбитражный суд {law_range}').bold = True
    # получаем объект форматирования
    fmt = p.paragraph_format
    # Форматируем:
    # отступы
    fmt.left_indent = Mm(70)
    fmt.right_indent = Mm(-10)
    # Межстрочный интервал:
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


    p = doc.add_paragraph()
    p.add_run(law_range_address).bold = True
    # получаем объект форматирования
    fmt = p.paragraph_format
    # Форматируем:
    # отступы
    fmt.left_indent = Mm(70)
    fmt.right_indent = Mm(-10)
    # Межстрочный интервал:
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


    p = doc.add_paragraph()
    p.add_run(f'А{law_number_first}-{law_number_last}/{law_number_year}\n').bold = True
    # получаем объект форматирования
    fmt = p.paragraph_format
    # Форматируем:
    # отступы
    fmt.left_indent = Mm(70)
    fmt.right_indent = Mm(-10)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


    p = doc.add_paragraph()
    run=p.add_run('Должник')
    run.underline = True
    run.bold = True
    run.italic=True
    # получаем объект форматирования
    fmt = p.paragraph_format
    # Форматируем:
    # отступы
    fmt.left_indent = Mm(70)
    fmt.right_indent = Mm(-10)
    # Межстрочный интервал:
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


    p = doc.add_paragraph()
    a=p.add_run(physics)
    # получаем объект форматирования
    fmt = p.paragraph_format
    # Форматируем:
    # отступы
    fmt.left_indent = Mm(70)
    fmt.right_indent = Mm(-10)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


    p = doc.add_paragraph()
    run=p.add_run(f'{physics_address}\n')
    run.bold=True
    # получаем объект форматирования
    fmt = p.paragraph_format
    # Форматируем:
    # отступы
    fmt.left_indent = Mm(70)
    fmt.right_indent = Mm(-10)
    # Межстрочный интервал:
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    p = doc.add_paragraph()
    run=p.add_run('Кредитор:')
    run.bold=True
    run.italic=True
    # получаем объект форматирования
    fmt = p.paragraph_format
    # Форматируем:
    # отступы
    fmt.left_indent = Mm(70)
    fmt.right_indent = Mm(-10)
    # Межстрочный интервал:
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


    p = doc.add_paragraph()
    run=p.add_run(creditor)
    run.bold=True
    run.italic=True
    # получаем объект форматирования
    fmt = p.paragraph_format
    # Форматируем:
    # отступы
    fmt.left_indent = Mm(70)
    fmt.right_indent = Mm(-10)
    # Межстрочный интервал:
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


    p = doc.add_paragraph()
    run=p.add_run(creditor_address)
    # получаем объект форматирования
    fmt = p.paragraph_format
    # Форматируем:
    # отступы
    fmt.left_indent = Mm(70)
    fmt.right_indent = Mm(-10)
    # Межстрочный интервал:
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    p = doc.add_paragraph()
    run=p.add_run('Финансовый управляющий Кондакова Ирина Николаевна')
    run.bold=True
    run.underline = True
    run.italic=True
    run=p.add_run(', ИНН 550714197767 ')
    run=p.add_run('член Ассоциации «Национальная организация арбитражных управляющих»,')
    run.bold=True
    # Межстрочный интервал:
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # получаем объект форматирования
    fmt = p.paragraph_format
    # Форматируем:
    # отступы
    fmt.left_indent = Mm(70)
    fmt.right_indent = Mm(-10)
    # Межстрочный интервал:
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    p = doc.add_paragraph()
    run=p.add_run('Адрес: 644081, г. Омск, ул. Фугенфирова, д. 2, кв. 58\n')
    # получаем объект форматирования
    fmt = p.paragraph_format
    # Форматируем:
    # отступы
    fmt.left_indent = Mm(70)
    fmt.right_indent = Mm(-10)
    # Межстрочный интервал:
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    #==============================Тело============================================#

    p = doc.add_paragraph()
    run=p.add_run('ОТЗЫВ')
    run.bold = True
    # Размер текста
    run.font.size=Pt(14)
    # Положение текста посередине:
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    p = doc.add_paragraph()
    run=p.add_run('на заявление кредитора о включении в реестр требований кредиторов')
    run.bold = True
    run.font.size=Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    p = doc.add_paragraph()
    run=p.add_run('От')
    run=p.add_run(f' {creditor} ({creditor_address}) ')
    run.bold=True
    run=p.add_run(f'в Арбитражный суд {law_range} поступило заявление о включении в реестр требований кредиторов должника {physics} ({creditor_address}), {summ} руб. Не возражаю против включения требования кредитора в реестр требований кредиторов должника\n')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    p = doc.add_paragraph()
    a=p.add_run('Составу суда доверяю, отводов не имею')
    # получаем объект форматирования
    fmt = p.paragraph_format
    # Форматируем:
    # отступы
    fmt.left_indent = Mm(15)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    p = doc.add_paragraph()
    run=p.add_run('Приложение:\n1.	Уведомление о получении требования кредитора на сайте «Федресурс».\n2.	Копия квитанции о направлении отзыва кредитору.')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    p = doc.add_paragraph()
    run=p.add_run('Финансовый управляющий')
    run.add_picture('photo.jpg',Mm(24),Mm(15))
    p.add_run('Кондакова Ирина Николаевна')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


    physics_familian=physics.split(' ')[0]
    doc.save(f'Отзыв {physics_familian} {creditor}.docx')

#============================Получение данных=================================#
law_range=input('Арбитражный суд')
law_range_address=input('Адрес:')
law_number_first=input('Номер после А: ')
law_number_last=input('Номер после -: ')
law_number_year=input('Год: ')
physics=input('ФИО Должника: ')
physics_address=input('Его место жительства: ')
creditor=input('Кредитор: ')
creditor_address=input('Его адрес: ')
summ=input('Сумма долга: ')

create_file(law_range,law_range_address,law_number_first,law_number_last,law_number_year,physics,physics_address,creditor,creditor_address,summ)
